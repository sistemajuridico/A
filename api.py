import os
import io
import json
import shutil
import time
import uuid
import unicodedata
import traceback
from fastapi import FastAPI, UploadFile, File, Form, BackgroundTasks, Request
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse
from pydantic import BaseModel
from typing import List, Optional
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from google import genai
from google.genai import types

app = FastAPI()

MAX_UPLOAD_SIZE = 300 * 1024 * 1024  

@app.middleware("http")
async def limit_upload_size(request: Request, call_next):
    if request.method == "POST":
        content_length = request.headers.get("content-length")
        if content_length and int(content_length) > MAX_UPLOAD_SIZE:
            return JSONResponse(
                {"erro": "A soma total dos ficheiros excede 300MB. Envie menos volumes por vez."},
                status_code=413
            )
    return await call_next(request)

TASKS = {}

@app.get("/")
def serve_index():
    return FileResponse("index.html")

@app.head("/")
def ping_render():
    """Responde ao 'Health Check' do Render para ele não matar o servidor"""
    return JSONResponse(content={"status": "alive"})

# --- FATIADOR AUTOMÁTICO DE PDFS GIGANTES ---
def dividir_pdf_se_necessario(caminho_pdf, max_paginas=900):
    """Garante que a Google não bloqueia PDFs com mais de 1000 páginas"""
    import PyPDF2
    arquivos_divididos = []
    try:
        with open(caminho_pdf, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            total_paginas = len(reader.pages)

            if total_paginas <= max_paginas:
                return [caminho_pdf] # Retorna o original se for pequeno

            print(f"A dividir PDF gigante de {total_paginas} páginas...")
            for i in range(0, total_paginas, max_paginas):
                writer = PyPDF2.PdfWriter()
                fim = min(i + max_paginas, total_paginas)
                for j in range(i, fim):
                    writer.add_page(reader.pages[j])

                novo_nome = f"{caminho_pdf}_parte_{i//max_paginas + 1}.pdf"
                with open(novo_nome, 'wb') as out_f:
                    writer.write(out_f)
                arquivos_divididos.append(novo_nome)
                
        return arquivos_divididos
    except Exception as e:
        print(f"Erro ao tentar dividir o PDF: {e}")
        return [caminho_pdf] # Em caso de erro, tenta enviar o original

def comprimir_video(input_path, output_path):
    try:
        from moviepy.editor import VideoFileClip
        with VideoFileClip(input_path) as video:
            video_redimensionado = video.resize(height=480)
            video_redimensionado.write_videofile(output_path, fps=15, codec="libx264", audio_codec="aac", logger=None)
        return True
    except Exception as e:
        return False

def comprimir_audio(input_path, output_path):
    try:
        from pydub import AudioSegment
        audio = AudioSegment.from_file(input_path)
        audio = audio.set_channels(1).set_frame_rate(16000)
        audio.export(output_path, format="mp3", bitrate="64k")
        return True
    except Exception as e:
        return False

def processar_background(task_id: str, fatos: str, area: str, mag: str, trib: str, arquivos_brutos: list):
    arquivos_para_gemini = []
    try:
        area = unicodedata.normalize('NFC', area)
        mag = unicodedata.normalize('NFC', mag)
        trib = unicodedata.normalize('NFC', trib)
        fatos = unicodedata.normalize('NFC', fatos)

        for temp_input, ext, safe_name in arquivos_brutos:
            if ext == "pdf": 
                # MÁGICA ACONTECE AQUI: Fatiamos o PDF antes de chegar no Google
                pedacos = dividir_pdf_se_necessario(temp_input)
                for pedaco in pedacos:
                    arquivos_para_gemini.append((pedaco, "application/pdf"))
                    
            elif ext in ["mp4", "mpeg", "mov", "avi"]:
                temp_output = f"comp_{safe_name}"
                if comprimir_video(temp_input, temp_output):
                    arquivos_para_gemini.append((temp_output, "video/mp4"))
                else:
                    arquivos_para_gemini.append((temp_input, "video/mp4"))
                    
            elif ext in ["mp3", "wav", "m4a", "ogg"]:
                temp_output = f"comp_{safe_name}"
                if comprimir_audio(temp_input, temp_output):
                    arquivos_para_gemini.append((temp_output, "audio/mp3"))
                else:
                    arquivos_para_gemini.append((temp_input, "audio/mp3"))

        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            TASKS[task_id] = {"status": "error", "erro": "Chave API não configurada."}
            return

        client = genai.Client(api_key=api_key)
        conteudos_multimais = []

        for target_file, mime in arquivos_para_gemini:
            gemini_file = client.files.upload(file=target_file, config={'mime_type': mime})
            while True:
                f_info = client.files.get(name=gemini_file.name)
                state_str = str(f_info.state).upper()
                if "FAILED" in state_str:
                    raise Exception("A IA falhou ao processar o ficheiro nos servidores da Google.")
                if "ACTIVE" in state_str:
                    break
                time.sleep(3)
            
            conteudos_multimais.append(
                types.Part.from_uri(file_uri=f_info.uri, mime_type=mime)
            )

        # --- NOVA ARQUITETURA MENTAL (CHAIN OF THOUGHT) ---
        instrucoes = f"""
        Você é o M.A | JUS IA EXPERIENCE, um Advogado de Elite e Estrategista Processual. Especialidade: {area}.
        
        MUDANÇA DE ARQUITETURA MENTAL (CHAIN OF THOUGHT):
        Você não é um mero gerador de texto. O seu processo mental OBRIGATÓRIO deve seguir esta ordem exata:
        
        PASSO 1: DIAGNÓSTICO. Analise os fatos e documentos. Crie um resumo estratégico focado no "calcanhar de Aquiles" do caso.
        PASSO 2: MODO COMBATE (FRAQUEZAS). Identifique meticulosamente todas as vulnerabilidades, contradições, prescrições ou teses fracas da contraparte.
        PASSO 3: ARSENAL JURÍDICO. Reúna a base legal, a jurisprudência e a doutrina que atacam EXATAMENTE as fraquezas encontradas no Passo 2.
        PASSO 4: A PEÇA DE RESISTÊNCIA. APENAS APÓS concluir os passos 1, 2 e 3, você irá redigir a PEÇA PROCESSUAL. A peça não pode ser genérica. Ela DEVE OBRIGATORIAMENTE usar como munição todas as fraquezas e jurisprudências que você levantou nos passos anteriores para aniquilar a tese adversária.
        
        REGRA CRÍTICA: É ESTRITAMENTE PROIBIDO copiar ou transcrever as petições antigas do PDF. O PDF é apenas o histórico. O seu trabalho é redigir uma PEÇA NOVA, INÉDITA e com a DATA ATUAL, rebatendo o que está no PDF e usando os Fatos Novos.
        
        ATENÇÃO MÁXIMA PARA A PEÇA PROCESSUAL: No campo 'peca_processual', você é PROIBIDO de resumir. Você DEVE redigir a NOVA PETIÇÃO COMPLETA, EXTENSA e PRONTA PARA PROTOCOLO (Endereçamento, Qualificação, Dos Fatos, Do Direito, Dos Pedidos e Fecho).

        RETORNE ESTRITAMENTE EM JSON COM ESTA ESTRUTURA E ORDEM MENTAL:
        {{
            "resumo_estrategico": "A sua tese principal (Passo 1)",
            "vulnerabilidades_contraparte": ["Fraqueza 1...", "Fraqueza 2... (Passo 2)"],
            "base_legal": ["Artigos essenciais... (Passo 3)"],
            "jurisprudencia": ["Julgados que atacam a contraparte... (Passo 3)"],
            "doutrina": ["Entendimento doutrinário... (Passo 3)"],
            "jurimetria": "Tendências do juízo",
            "timeline": ["Cronologia..."],
            "checklist": ["Providências práticas..."],
            "resumo_cliente": "Explicação leiga...",
            "peca_processual": "A NOVA PETIÇÃO COMPLETA E EXTENSA, CONSTRUÍDA COM AS ARMAS DOS PASSOS ANTERIORES (Passo 4)..."
        }}
        """
        
        prompt_partes = []
        prompt_partes.extend(conteudos_multimais)
        prompt_partes.append(f"{instrucoes}\n\nFATOS (CRIAR PEÇA NOVA COM BASE NISTO):\n{fatos}")

        # --- A VACINA DOS FILTROS (Permite analisar crimes, litígios e afins sem a Google bloquear) ---
        filtros_seguranca = [
            types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
            types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
        ]

        if len(conteudos_multimais) > 0:
            config_ia = types.GenerateContentConfig(
                temperature=0.1,
                response_mime_type="application/json",
                safety_settings=filtros_seguranca
            )
        else:
            config_ia = types.GenerateContentConfig(
                temperature=0.1, 
                response_mime_type="application/json",
                tools=[{"googleSearch": {}}],
                safety_settings=filtros_seguranca
            )

        response = client.models.generate_content(
            model='gemini-2.5-flash', 
            contents=prompt_partes,
            config=config_ia
        )

        # --- A REDE DE PROTEÇÃO CONTRA O NONETYPE ---
        if getattr(response, 'text', None) is None:
            motivo = "A Google bloqueou a resposta silenciosamente."
            if hasattr(response, 'candidates') and response.candidates and hasattr(response.candidates[0], 'finish_reason'):
                motivo = f"A IA recusou-se a gerar o texto. Motivo oficial da Google: {response.candidates[0].finish_reason}"
            raise Exception(motivo)

        texto_puro = response.text.strip()
        if texto_puro.startswith("```json"):
            texto_puro = texto_puro.replace("```json", "", 1)
        if texto_puro.endswith("```"):
            texto_puro = texto_puro.rsplit("```", 1)[0]
            
        TASKS[task_id] = {"status": "done", "resultado": json.loads(texto_puro.strip())}

    except Exception as e:
        print("=== DETALHE DO ERRO ===")
        print(traceback.format_exc())
        erro_seguro = str(e).encode('ascii', 'replace').decode('ascii')
        
        if "INVALID_ARGUMENT" in erro_seguro.upper() or "400" in erro_seguro:
            mensagem = "Erro 400 (Invalid Argument): A formatação do documento não é suportada ou o ficheiro está protegido por palavra-passe."
        else:
            mensagem = f"Erro na IA: {erro_seguro}"
            
        TASKS[task_id] = {"status": "error", "erro": mensagem}
    finally:
        # Limpeza total para não acumular "lixo" no servidor Render
        for f, m in arquivos_para_gemini:
            if os.path.exists(f): os.remove(f)
        for temp_input, ext, safe_name in arquivos_brutos:
            if os.path.exists(temp_input): os.remove(temp_input)

@app.post("/analisar")
async def analisar_caso(
    background_tasks: BackgroundTasks,
    fatos_do_caso: str = Form(default=""),
    area_direito: str = Form(default=""),
    magistrado: str = Form(default=""),
    tribunal: str = Form(default=""),
    arquivos: Optional[List[UploadFile]] = File(default=[])
):
    fatos_limpos = unicodedata.normalize('NFC', fatos_do_caso) if fatos_do_caso else ""
    mag_limpo = unicodedata.normalize('NFC', magistrado) if magistrado else ""
    trib_limpo = unicodedata.normalize('NFC', tribunal) if tribunal else ""

    if not fatos_limpos or len(fatos_limpos.strip()) < 5:
        return JSONResponse(content={"erro": "Descreva os fatos."}, status_code=400)

    task_id = str(uuid.uuid4())
    TASKS[task_id] = {"status": "processing"}

    arquivos_brutos = []
    try:
        if arquivos:
            for arquivo in arquivos:
                if not arquivo.filename: continue
                ext = arquivo.filename.lower().split('.')[-1]
                safe_name = f"doc_{uuid.uuid4().hex}.{ext}"
                temp_input = f"temp_in_{safe_name}"
                
                with open(temp_input, "wb") as buffer:
                    while True:
                        chunk = await arquivo.read(1024 * 1024)
                        if not chunk:
                            break
                        buffer.write(chunk)
                
                arquivos_brutos.append((temp_input, ext, safe_name))
                
    except Exception as e:
        return JSONResponse(content={"erro": "Erro de codificação ao salvar o ficheiro."}, status_code=500)

    background_tasks.add_task(processar_background, task_id, fatos_limpos, area_direito, mag_limpo, trib_limpo, arquivos_brutos)
    
    return JSONResponse(content={"task_id": task_id})

@app.get("/status/{task_id}")
def check_status(task_id: str):
    task = TASKS.get(task_id)
    if not task:
        return JSONResponse(content={"status": "error", "erro": "Tarefa perdida pelo servidor (Render reiniciou a máquina). Tente novamente."})
    return JSONResponse(content=task)

class DadosPeca(BaseModel):
    texto_peca: str
    advogado_nome: Optional[str] = ""
    advogado_oab: Optional[str] = ""
    advogado_endereco: Optional[str] = ""

@app.post("/gerar_docx")
def gerar_docx(dados: DadosPeca):
    try:
        doc = docx.Document()
        for s in doc.sections:
            s.top_margin, s.bottom_margin, s.left_margin, s.right_margin = Cm(3), Cm(2), Cm(3), Cm(2)

        if dados.advogado_nome:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_h = p.add_run(f"{str(dados.advogado_nome).upper()}\nOAB: {dados.advogado_oab if dados.advogado_oab else '---'}\n{dados.advogado_endereco if dados.advogado_endereco else ''}")
            run_h.font.size, run_h.font.name, run_h.italic = Pt(10), 'Times New Roman', True

        for linha in dados.texto_peca.split('\n'):
            if linha.strip():
                para = doc.add_paragraph(linha.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                para.paragraph_format.first_line_indent = Cm(2.0)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=MA_Elite.docx"})
    except Exception as e:
        return JSONResponse(content={"erro": "Erro na geração do ficheiro Word."}, status_code=500)
