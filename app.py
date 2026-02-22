import streamlit as st
import json
import os
import io
import PyPDF2
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types

# 1. CONFIGURAÇÃO DA PÁGINA E ESTILIZAÇÃO
st.set_page_config(
    page_title="M.A | Inteligência Jurídica",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS Customizado para LIGHT MODE CORPORATIVO PREMIUM
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600&family=Playfair+Display:wght@600;700&display=swap');

    /* Fundos Gerais */
    .stApp, [data-testid="stAppViewContainer"] { 
        background-color: #f8fafc; /* Fundo cinza super claro (gelo), mais elegante que o branco puro */
        font-family: 'Inter', sans-serif;
    }
    [data-testid="stHeader"] { background-color: transparent; }
    [data-testid="stSidebar"] { 
        background-color: #ffffff !important; 
        border-right: 1px solid #e2e8f0;
    }
    
    /* Tipografia de Títulos (Estilo Editorial/Jurídico) */
    h1, h2, h3 { 
        color: #0f172a !important; /* Azul marinho quase preto */
        font-family: 'Playfair Display', serif !important; 
        font-weight: 700; 
        letter-spacing: -0.5px;
    }
    p, label, .stMarkdown, span { color: #334155 !important; font-family: 'Inter', sans-serif; }
    
    /* Botão Principal - Destaque vibrante */
    .stButton>button {
        background: linear-gradient(135deg, #1e3a8a, #2563eb);
        color: white !important;
        border-radius: 8px;
        padding: 14px 28px;
        font-weight: 600;
        font-size: 1.05rem;
        letter-spacing: 0.5px;
        border: none;
        box-shadow: 0 4px 14px rgba(37, 99, 235, 0.3);
        transition: all 0.3s ease;
        width: 100%;
    }
    .stButton>button:hover { 
        background: linear-gradient(135deg, #2563eb, #3b82f6); 
        box-shadow: 0 6px 20px rgba(37, 99, 235, 0.4);
        transform: translateY(-2px);
    }
    
    /* Caixa de Tese Principal (Branca com Sombra Elegante) */
    .estilo-caixa {
        background: #ffffff;
        padding: 30px;
        border-radius: 12px;
        box-shadow: 0 10px 25px rgba(0, 0, 0, 0.05); /* Sombra suave para "saltar" da tela */
        border: 1px solid #f1f5f9;
        border-left: 6px solid #2563eb; /* Faixa azul lateral */
        margin-bottom: 30px;
        color: #1e293b;
    }
    .estilo-caixa h3 { color: #1e3a8a !important; }
    
    /* Inputs e Áreas de Texto */
    .stTextArea textarea, .stTextInput input, .stSelectbox div[data-baseweb="select"] {
        background-color: #ffffff !important;
        color: #0f172a !important;
        border: 1px solid #cbd5e1 !important;
        border-radius: 8px;
        font-size: 1rem;
        box-shadow: 0 1px 2px rgba(0,0,0,0.02) inset;
        transition: border-color 0.3s ease, box-shadow 0.3s ease;
    }
    .stTextArea textarea:focus, .stTextInput input:focus {
        border-color: #2563eb !important;
        box-shadow: 0 0 0 2px rgba(37, 99, 235, 0.2) !important;
    }
    
    /* Estilo das Abas (Tabs) Refinado */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background-color: transparent;
        border-bottom: 1px solid #cbd5e1;
        padding-bottom: 0;
    }
    .stTabs [data-baseweb="tab"] {
        height: 45px;
        white-space: pre-wrap;
        background-color: transparent;
        border-radius: 6px 6px 0px 0px;
        border: none;
        color: #64748b !important;
        font-weight: 500;
        font-family: 'Inter', sans-serif;
        transition: color 0.2s ease;
    }
    .stTabs [aria-selected="true"] {
        background-color: transparent !important;
        color: #2563eb !important;
        border-bottom: 3px solid #2563eb !important;
    }
    
    /* Uploader de arquivo - Minimalista */
    [data-testid="stFileUploadDropzone"] {
        background-color: #f1f5f9 !important;
        border: 2px dashed #94a3b8 !important;
        border-radius: 12px;
        padding: 30px;
        transition: all 0.3s ease;
    }
    [data-testid="stFileUploadDropzone"]:hover {
        border-color: #2563eb !important;
        background-color: #e2e8f0 !important;
    }

    /* Cards de Resultados (Base Legal, Jurisp., Doutrina) - Levemente Coloridos */
    .result-card {
        padding: 20px;
        border-radius: 8px;
        margin-bottom: 15px;
        border: 1px solid #e2e8f0;
        color: #1e293b;
        font-size: 0.98rem;
        line-height: 1.6;
        box-shadow: 0 2px 5px rgba(0,0,0,0.02);
    }
    .card-legal { border-left: 4px solid #3b82f6; background-color: #eff6ff; } /* Azul clarinho */
    .card-juris { border-left: 4px solid #f59e0b; background-color: #fffbeb; } /* Amarelo clarinho */
    .card-doutrina { border-left: 4px solid #10b981; background-color: #ecfdf5; } /* Verde clarinho */

    hr { border-color: #e2e8f0; }
    
    /* Ajuste de ícones e textos secundários */
    .st-emotion-cache-1wmy9hl { color: #64748b; }
    </style>
""", unsafe_allow_html=True)

# 2. GERENCIAMENTO DE CONFIGURAÇÕES
ARQUIVO_CONFIG = "config_ma.json"

def carregar_config():
    if os.path.exists(ARQUIVO_CONFIG):
        with open(ARQUIVO_CONFIG, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"api_key": "", "advogado_nome": "", "advogado_oab": "", "advogado_endereco": ""}

def salvar_config(dados):
    with open(ARQUIVO_CONFIG, "w", encoding="utf-8") as f:
        json.dump(dados, f, indent=4)

if "config" not in st.session_state:
    st.session_state.config = carregar_config()

def extrair_texto_pdf(arquivo_pdf):
    texto = ""
    try:
        leitor = PyPDF2.PdfReader(arquivo_pdf)
        for pagina in leitor.pages:
            texto += pagina.extract_text() + "\n"
    except Exception as e:
        st.error(f"Erro ao ler o PDF: {e}")
    return texto

def gerar_docx(texto_peca, dados_advogado):
    doc = docx.Document()
    
    if dados_advogado['advogado_nome']:
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_nome = header.add_run(dados_advogado['advogado_nome'].upper())
        run_nome.bold = True
        run_nome.font.size = Pt(14)
        
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_info = info.add_run(f"OAB: {dados_advogado['advogado_oab']} | {dados_advogado['advogado_endereco']}")
        run_info.font.size = Pt(9)
        run_info.italic = True
        
        doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("\n")

    estilo = doc.styles['Normal']
    fonte = estilo.font
    fonte.name = 'Arial'
    fonte.size = Pt(12)
    
    paragrafos = texto_peca.split('\n')
    for p in paragrafos:
        if p.strip():
            para = doc.add_paragraph(p.strip())
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 3. FUNÇÃO DO MOTOR DE IA
def realizar_pesquisa_processual(fatos_do_caso: str, texto_documentos: str, area_direito: str, api_key: str) -> dict:
    try:
        cliente = genai.Client(api_key=api_key)
        
        instrucoes_sistema = f"""
        Você é um advogado sênior, jurista renomado e pesquisador especialista em {area_direito} no Brasil.
        Sua missão é atuar na ETAPA 1 de um caso: A Pesquisa e Análise Processual Estratégica.
        
        DIRETRIZES OBRIGATÓRIAS:
        1. Responda ESTRITAMENTE em Português do Brasil (PT-BR).
        2. Utilize vernáculo jurídico adequado, formal e profissional.
        3. Você TEM ACESSO À INTERNET através do Google Search. É OBRIGATÓRIO buscar jurisprudência real, atualizada e verídica. NÃO invente números.
        
        Responda EXCLUSIVAMENTE em formato JSON com a seguinte estrutura exata:
        {{
            "resumo_estrategico": "texto do resumo claro, direto e persuasivo",
            "base_legal": ["Artigo X da Lei Y: Explicação", "Artigo Z..."],
            "jurisprudencia": ["Tribunal (ex: STJ) - Tema/Súmula: Explicação", "TJSP..."],
            "doutrina": ["Nome do Autor: Resumo do entendimento", "Outro Autor..."],
            "peca_processual": "Texto COMPLETO da peça processual com quebras de linha (\\n)."
        }}
        """

        prompt_completo = f"{instrucoes_sistema}\n\n"
        if texto_documentos.strip():
            prompt_completo += f"--- DOCUMENTOS DO PROCESSO ---\n{texto_documentos}\n--- FIM ---\n\n"
        prompt_completo += f"FATOS DO CASO:\n{fatos_do_caso}"

        resposta = cliente.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt_completo,
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                temperature=0.2,
                tools=[{"google_search": {}}]
            )
        )
        return json.loads(resposta.text)

    except Exception as e:
        return {"erro": str(e)}

# 4. INTERFACE VISUAL PRINCIPAL
st.markdown("<h1 style='text-align: center; margin-bottom: 0.5rem; color: #0f172a;'>M.A <span style='color: #2563eb;'>|</span> Inteligência Jurídica</h1>", unsafe_allow_html=True)
st.markdown("<p style='text-align: center; font-size: 1.1rem; color: #64748b !important; margin-bottom: 3rem;'>Sistema avançado de apoio à decisão e pesquisa jurisprudencial</p>", unsafe_allow_html=True)

# --- BARRA LATERAL ---
with st.sidebar:
    st.markdown("<h3 style='margin-bottom: 20px; color: #0f172a;'>⚙️ Painel de Controle</h3>", unsafe_allow_html=True)
    
    with st.expander("🔑 Credenciais da IA", expanded=True):
        api_key_input = st.text_input("Chave API (Google Gemini):", value=st.session_state.config["api_key"], type="password")
    
    with st.expander("👤 Dados da Assinatura (Peça)"):
        nome_adv = st.text_input("Nome do Advogado(a):", value=st.session_state.config["advogado_nome"])
        oab_adv = st.text_input("OAB:", value=st.session_state.config["advogado_oab"])
        end_adv = st.text_area("Endereço/Contato:", value=st.session_state.config["advogado_endereco"], height=100)
    
    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("💾 Salvar Configurações"):
        st.session_state.config = {"api_key": api_key_input, "advogado_nome": nome_adv, "advogado_oab": oab_adv, "advogado_endereco": end_adv}
        salvar_config(st.session_state.config)
        st.success("Configurações salvas com sucesso!")

# --- ÁREA DE INPUT (Layout em Colunas) ---
st.markdown("<h3 style='color: #0f172a;'>📋 Configuração do Caso</h3>", unsafe_allow_html=True)
st.markdown("<br>", unsafe_allow_html=True)

col_esq, col_dir = st.columns([1, 2], gap="large")

with col_esq:
    area_selecionada = st.selectbox(
        "Ramo do Direito Aplicável:",
        ["Direito Civil, Imobiliário e Consumidor", "Direito de Família e Sucessões", "Direito Penal e Processual Penal", "Direito Previdenciário", "Direito do Trabalho", "Direito Tributário e Empresarial"]
    )
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("<p style='font-weight: 600; margin-bottom: 5px; color: #334155;'>📄 Autos do Processo (Opcional)</p>", unsafe_allow_html=True)
    arquivos_anexados = st.file_uploader("Arraste PDFs iniciais, B.O. ou contratos aqui", type=["pdf"], accept_multiple_files=True, label_visibility="collapsed")

with col_dir:
    fatos_input = st.text_area(
        "📝 Relato Estratégico e Instruções:", 
        height=240, 
        placeholder="Descreva detalhadamente os fatos do caso. Ex:\n\n'Meu cliente sofreu um golpe via Pix. O banco recebedor da fraude não bloqueou a conta mesmo após o alerta MED. Quero uma ação indenizatória focada na súmula 479 do STJ...'"
    )

st.markdown("<br>", unsafe_allow_html=True)

# Botão Centralizado e Destaque
_, col_btn, _ = st.columns([1, 2, 1])
with col_btn:
    executar = st.button("⚖️ Executar Análise Jurídica Avançada")

st.markdown("<hr style='margin-top: 3rem; margin-bottom: 3rem;'>", unsafe_allow_html=True)

# --- EXECUÇÃO E RESULTADOS ---
if executar:
    if not st.session_state.config["api_key"]:
        st.error("⚠️ Atenção: Configure sua Chave da API no painel lateral esquerdo.")
    elif len(fatos_input.strip()) < 10 and not arquivos_anexados:
        st.warning("⚠️ Forneça um relato mínimo dos fatos ou anexe documentos para análise.")
    else:
        with st.spinner('🔍 Analisando doutrina, consultando tribunais e estruturando tese...'):
            texto_extraido = ""
            if arquivos_anexados:
                for arq in arquivos_anexados:
                    texto_extraido += f"\n--- {arq.name} ---\n{extrair_texto_pdf(arq)}"
            
            resultado = realizar_pesquisa_processual(fatos_input, texto_extraido, area_selecionada, st.session_state.config["api_key"])
            
            if "erro" in resultado:
                st.error(f"❌ Erro na comunicação com a IA: {resultado['erro']}")
            else:
                st.markdown("<h2 style='color: #0f172a;'>📊 Parecer Estratégico M.A</h2>", unsafe_allow_html=True)
                
                # Tese Principal destacada
                st.markdown(f"""
                <div class="estilo-caixa">
                    <h3 style='margin-top: 0; font-size: 1.4rem;'>📌 Tese Principal Formada</h3>
                    <p style='font-size: 1.1rem; line-height: 1.7;'>{resultado.get("resumo_estrategico", "")}</p>
                </div>
                """, unsafe_allow_html=True)
                
                # Sistema de Abas Redesenhado
                tab1, tab2, tab3 = st.tabs(["⚖️ Fundamentação Legal", "🏛️ Jurisprudência Consolidada", "📚 Entendimento Doutrinário"])
                
                with tab1:
                    st.markdown("<br>", unsafe_allow_html=True)
                    for item in resultado.get("base_legal", []):
                        st.markdown(f'<div class="result-card card-legal">📖 <strong>Dispositivo:</strong> {item}</div>', unsafe_allow_html=True)
                        
                with tab2:
                    st.markdown("<br>", unsafe_allow_html=True)
                    for item in resultado.get("jurisprudencia", []):
                        st.markdown(f'<div class="result-card card-juris">⚖️ <strong>Precedente:</strong> {item}</div>', unsafe_allow_html=True)
                        
                with tab3:
                    st.markdown("<br>", unsafe_allow_html=True)
                    for item in resultado.get("doutrina", []):
                        st.markdown(f'<div class="result-card card-doutrina">✍️ <strong>Doutrina:</strong> {item}</div>', unsafe_allow_html=True)
                        
                # Geração da Peça
                peca_texto = resultado.get("peca_processual", "")
                if peca_texto:
                    st.markdown("<br><br>", unsafe_allow_html=True)
                    st.markdown("<h3 style='color: #0f172a;'>📄 Minuta da Peça Processual Gerada</h3>", unsafe_allow_html=True)
                    st.text_area("Revisão Rápida da Peça (Editável):", peca_texto, height=400)
                    
                    docx_buffer = gerar_docx(peca_texto, st.session_state.config)
                    
                    _, col_down, _ = st.columns([1, 2, 1])
                    with col_down:
                        st.download_button(
                            label="⬇️ Exportar Peça em Microsoft Word (.docx)",
                            data=docx_buffer,
                            file_name="peca_processual_MA.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
